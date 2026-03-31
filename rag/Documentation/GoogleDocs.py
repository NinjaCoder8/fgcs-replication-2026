from docx import Document
from docx.shared import Pt

doc = Document()

doc.add_heading("Project Pitch: Structured Knowledge Extraction in MLOps", level=0)

sections = [
    ("Introduction", 
     "Research in Machine Learning Operations (MLOps) is extensive and fragmented. "
     "In this project, we started with approximately forty PhD-level research papers that collectively described two major elements:\n\n"
     "- Actors: Sixteen distinct roles involved in MLOps, such as Data Scientists, Solution Architects, and MLOps Engineers.\n"
     "- Activities: More than four hundred tasks distributed across the MLOps lifecycle."),

    ("Purpose of the Project", 
     "We designed and implemented a pipeline that:\n"
     "- Parses the papers, splits them into manageable chunks, and encodes them as embeddings.\n"
     "- Uses a FAISS index to enable fast semantic retrieval of relevant text.\n"
     "- Builds task-specific prompts that ensure all reasoning is based only on the academic content.\n"
     "- Produces structured outputs, including:\n"
     "  * Actor groupings into dimensions, with a proposed leading role for each.\n"
     "  * Classification of all activities into the five core MLOps dimensions: Plan, Data Engineering, Model Engineering, Software Engineering, and Operations.\n"
     "  * Grouped and standardized activity names to remove redundancy and ensure clarity."),

    ("Model Comparison",
     "A central objective of the project was to compare results across two different AI models:\n\n"
     "- OpenAI GPT-4o: A large, cloud-based model with roughly 200B parameters and a 128K token context window.\n"
     "- Mistral-7B-Instruct-v0.3: A smaller, open-source model with 7.25B parameters and a 32K token context window, running fully on local hardware.\n\n"
     "This dual approach allowed us to evaluate trade-offs between a powerful proprietary model and a more lightweight, private, and cost-free alternative. "
     "GPT-4o produced more consistent and detailed results, while Mistral demonstrated the feasibility of conducting the same analysis offline."),

    ("Outcome",
     "The pipeline successfully transformed a large body of dense academic work into a structured, ready-to-use knowledge base. "
     "At the same time, it provided insight into the comparative strengths and limitations of cloud-scale versus local open-source models for knowledge extraction tasks."),

    ("End-to-End Flow of the Pipeline",
     "The project implements a step-by-step pipeline that takes raw academic PDFs and transforms them into structured knowledge tables. Below is the detailed flow.\n\n"
     "1. PDF Parsing (parse_pdfs.py)\n"
     "   - Reads each paper page by page.\n"
     "   - Captures paper ID, page number, text content, and file path.\n"
     "   - Stores extracted content in parsed_pages.json.\n\n"
     "2. Text Chunking (chunk_text.py)\n"
     "   - Splits text into ~1000 character chunks with 200-character overlap.\n"
     "   - Preserves metadata (paper ID, page, chunk index).\n"
     "   - Saves to parsed_chunks.json.\n\n"
     "3. Embedding Generation (embedding.py)\n"
     "   - Encodes chunks using SentenceTransformer (bge-m3).\n"
     "   - Stores embeddings in a FAISS index (faiss_index.index).\n\n"
     "4. FAISS Index Validation (faiss_values.py)\n"
     "   - Verifies vector count and dimensionality.\n\n"
     "5. Querying (query.py)\n"
     "   - Converts queries into embeddings.\n"
     "   - Retrieves relevant chunks and assembles context-rich prompts.\n\n"
     "6. Reasoning Stages\n"
     "   - Two branches: GPT-4o and Mistral-7B.\n"
     "   - Both group actors, classify activities, and standardize names.\n\n"
     "7. Result Consolidation (multiple_sheets.py)\n"
     "   - Merges results into a structured Excel file with multiple sheets for actors, dimensions, and activities."),

    ("Summary",
     "The pipeline can be described as:\n\n"
     "Parse → Chunk → Embed → Index → Query → Reason (GPT-4o & Mistral-7B) → Consolidate\n\n"
     "It transforms raw academic PDFs into a structured dataset of actors and activities in MLOps, "
     "while also comparing proprietary and open-source reasoning models."),

    ("Conclusion",
     "Through this project, we tested the full pipeline using both OpenAI GPT-4o and the Mistral-7B-Instruct-v0.3 local model.\n\n"
     "- Accuracy and Quality of Results: GPT-4o consistently provided more accurate, consistent, and detailed outputs, benefiting from its scale and larger context window.\n"
     "- Practicality and Accessibility: Running Mistral-7B locally requires high-end GPUs, making it impractical for most users. GPT-4o runs seamlessly in the cloud.\n\n"
     "Overall Assessment:\n"
     "- Mistral-7B shows that open-source local reasoning is feasible and private.\n"
     "- GPT-4o is clearly superior in quality, scalability, and real-world applicability.\n\n"
     "Final Note: GPT-4o is the preferred choice for transforming dense academic content into structured knowledge bases, "
     "combining accuracy, accessibility, and practicality.")
]

for title, content in sections:
    doc.add_heading(title, level=1)
    for para in content.split("\n\n"):
        doc.add_paragraph(para)

output_path = r"C:\Users\(TECH MART)\Desktop\PhD-ai-agent\Documentation\MLOps_Project_Documentation.docx"
doc.save(output_path)

output_path
